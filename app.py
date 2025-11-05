from docx import Document
import torch
import torch.nn.functional as F
from torchvision import transforms
from PIL import Image
import yaml
from agent import RobustnessAgent
from torchvision import datasets, transforms
from torch.utils.data import DataLoader

device = 'cuda' if torch.cuda.is_available() else 'cpu'

model = torch.load("models/brain_tumor_full_model.pth", weights_only=False)
model.to(device).eval()

cfg = yaml.safe_load(open("config.yaml"))

test_transforms = transforms.Compose([
    transforms.Resize((224, 224)),
    transforms.ToTensor(),
    transforms.Normalize([0.5]*3, [0.5]*3)
])


def preprocess_image(img_path):
    img = Image.open(img_path).convert('RGB')
    return test_transforms(img).unsqueeze(0)


train_dataset = datasets.ImageFolder(
    "dataset/train", transform=test_transforms)
train_loader = DataLoader(train_dataset, batch_size=32, shuffle=True)
agent = RobustnessAgent(model, device, train_loader=train_loader, config=cfg)


def run_agentic_analysis(img_path):
    print("\n" + "=" * 80)
    print("AGENTIC AI ROBUSTNESS EVALUATION REPORT")
    print("=" * 80 + "\n")

    image_tensor = preprocess_image(img_path).to(device)

    with torch.no_grad():
        outputs = model(image_tensor)
        probs = F.softmax(outputs, dim=1)
        pred_class = torch.argmax(probs, dim=1).item()
        confidence = probs[0][pred_class].item()

    class_names = ["meningioma", "glioma", "pituitary"]
    pred_label = class_names[pred_class] if pred_class < len(
        class_names) else "Unknown"

    print("INITIAL CLASSIFICATION")
    print("-" * 80)
    print(f"Predicted Tumor Type : {pred_label}")
    print(f"Model Confidence     : {confidence * 100:.2f}%\n")

    print("AGENTIC DECISION PROCESS")
    print("-" * 80)

    decisions, record = agent.handle_uploaded_image(image_tensor)

    for i, (decision, reasoning) in enumerate(decisions, start=1):
        print(f"\nStep {i}: {decision.upper()}")
        print(f"Reasoning: {reasoning}")

    robust_val = record.get("robust_accuracy")
    final_val = record.get("post_retrain_robust")

    print("\nROBUSTNESS EVALUATION SUMMARY")
    print("-" * 80)
    print(f"Robust Accuracy (Before Retraining): "
          f"{robust_val:.2f}%" if robust_val else "Robust Accuracy (Before Retraining): N/A")
    print(f"Robust Accuracy (After Retraining) : "
          f"{final_val:.2f}%" if final_val else "Robust Accuracy (After Retraining) : N/A")

    print("\n" + "=" * 80)
    print("AGENTIC AI ANALYSIS COMPLETE")
    print("=" * 80 + "\n")

    return {
        "initial_class": pred_label,
        "confidence": confidence,
        "decisions": decisions,
        "robust_before": robust_val,
        "robust_after": final_val
    }


if __name__ == "__main__":
    img_path = "uploads/uploaded_image.png"
    results = run_agentic_analysis(img_path)
    print("\n=== FINAL SUMMARY ===")
    print(results)


# --- Generate DOCX Report ---


def generate_report_docx(results, output_path="Agentic_AI_Robustness_Report.docx"):
    document = Document()
    document.add_heading("AGENTIC AI ROBUSTNESS EVALUATION REPORT", level=1)

    document.add_heading("1. Initial Classification", level=2)
    document.add_paragraph(f"Tumor Type: {results['initial_class']}")
    document.add_paragraph(
        f"Model Confidence: {results['confidence']*100:.2f}%")

    document.add_heading("2. Agentic Decision Process", level=2)
    for i, (decision, reasoning) in enumerate(results['decisions'], start=1):
        document.add_paragraph(f"Step {i}: {decision}", style='List Number')
        document.add_paragraph(f"Reasoning: {reasoning}")

    document.add_heading("3. Robustness Evaluation Summary", level=2)
    before = (
        f"{results['robust_before']:.2f}%" if results['robust_before'] else "N/A")
    after = (
        f"{results['robust_after']:.2f}%" if results['robust_after'] else "N/A")
    document.add_paragraph(f"Robust Accuracy (Before Retraining): {before}")
    document.add_paragraph(f"Robust Accuracy (After Retraining): {after}")

    document.add_paragraph(
        "\nReport generated automatically by Agentic AI System.")
    document.save(output_path)
    print(f"\nDOCX report generated successfully: {output_path}")


# --- Call the report generator ---
generate_report_docx(results)
